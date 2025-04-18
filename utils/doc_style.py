from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches, Pt
from docx import Document
import io
from lib.htmldocx import HtmlToDocx
import logging

logger = logging.getLogger(__name__)

def apply_font_style(element, font_name, font_size):
    if hasattr(element, 'font'):
        element.font.name = font_name

def apply_font_to_elements(elements, font_name, font_size):
    for element in elements:
        apply_font_style(element, font_name, font_size)
        if hasattr(element, 'runs'):
            for run in element.runs:
                apply_font_style(run, font_name, font_size)

def get_subdoc(doc, raw_html, headers, base_url):
    temp_doc = Document()
    temp_parser = HtmlToDocx()

    if raw_html is not None:
        temp_parser.add_html_to_document(raw_html, temp_doc, headers, base_url)

        text_width = temp_doc.sections[0].page_width - temp_doc.sections[0].left_margin - temp_doc.sections[0].right_margin

        apply_font_to_elements(temp_doc.element.body, 'Calibri', 16)

        for paragraph in temp_doc.paragraphs:
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.left_indent = Pt(45)
            paragraph.paragraph_format.right_indent = Inches(1)
            paragraph.paragraph_format.top_indent = Inches(1)
            paragraph.paragraph_format.line_spacing = 1.5

        for style_name in ['Normal', 'List Bullet']:
            if style_name in temp_doc.styles:
                temp_doc.styles[style_name].font.name = 'Calibri'

        subdoc_tmp = io.BytesIO()
        temp_doc.save(subdoc_tmp)
        subdoc_tmp.seek(0)
        sub_docxtpl = DocxTemplate(subdoc_tmp)

        sub_docxtpl.save(subdoc_tmp)
        subdoc_tmp.seek(0)
        return doc.new_subdoc(subdoc_tmp)

def main_doc_style(doc):
    try:
        font = doc.styles['List Bullet'].font if 'List Bullet' in doc.styles else doc.styles['Normal'].font
        font.name = 'Calibri'
        font.size = Pt(16)

        # Use section[0] safely
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        for table in doc.tables:
            for row in table.rows:
                if all(cell.text.strip() == "" for cell in row.cells):
                    row_element = row._element
                    row_element.getparent().remove(row_element)
                else:
                    row.height = None
                    row.height_rule = None
        return doc

    except Exception as e:
        logger.error(f"main_doc_style error: {e}")
        return doc
   
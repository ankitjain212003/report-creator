import re, argparse
import io, os
import urllib.request
from urllib.parse import urlparse
from html.parser import HTMLParser
import requests
from requests.packages.urllib3.exceptions import InsecureRequestWarning
import docx, docx.table
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup

# values in inches
INDENT = 0.25
LIST_INDENT = 0.5
MAX_INDENT = 5.5  # To stop indents going off the page

# Style to use with tables. By default no style is used.
DEFAULT_TABLE_STYLE = None

# Style to use with paragraphs. By default no style is used.
DEFAULT_PARAGRAPH_STYLE = None


def get_filename_from_url(url):
    return os.path.basename(urlparse(url).path)


def is_url(url):
    """
    Not to be used for actually validating a url, but in our use case we only 
    care if it's a url or a file path, and they're pretty distinguishable
    """
    parts = urlparse(url)
    result = all([parts.scheme, parts.netloc, parts.path])
    print(f"is_url('{url}') -> {result}")
    return result


def fetch_image(url, headers, base_url):
    """
    Attempts to fetch an image from a url (absolute or relative).
    Returns a BytesIO on success, else None.
    """
    requests.packages.urllib3.disable_warnings(InsecureRequestWarning)

    # Determine full URL
    if url.startswith('http://') or url.startswith('https://'):
        full_url = url
    else:
        full_url = f"{base_url.rstrip('/')}/{url.lstrip('/')}"

    print("⏳ Fetching image from:", full_url)
    response = requests.get(
        full_url,
        headers={"Authorization": f"Bearer {headers}"},
        verify=False
    )
    if response.status_code == 200:
        return io.BytesIO(response.content)
    print("⚠️ Image fetch failed:", response.status_code, response.text)
    return None


def remove_last_occurence(ls, x):
    ls.pop(len(ls) - ls[::-1].index(x) - 1)


def remove_whitespace(string, leading=False, trailing=False):
    """Remove white space from a string."""
    if leading:
        string = re.sub(r'^\s*\n+\s*', '', string)
    if trailing:
        string = re.sub(r'\s*\n+\s*$', '', string)
    string = re.sub(r'\s*\n\s*', ' ', string)
    return re.sub(r'\s+', ' ', string)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


font_styles = {
    'b': 'bold', 'strong': 'bold', 'em': 'italic', 'i': 'italic',
    'u': 'underline', 's': 'strike', 'sup': 'superscript', 'sub': 'subscript',
    'th': 'bold',
}

font_names = {
    'code': 'Courier', 'pre': 'Courier',
}

styles = {
    'LIST_BULLET': 'List Bullet',
    'LIST_NUMBER': 'List Number',
}


class HtmlToDocx(HTMLParser):

    def __init__(self):
        super().__init__()
        self.options = {
            'fix-html': True,
            'images': True,
            'tables': True,
            'styles': True,
        }
        self.table_row_selectors = [
            'table > tr', 'table > thead > tr', 'table > tbody > tr', 'table > tfoot > tr'
        ]
        self.table_style = DEFAULT_TABLE_STYLE
        self.paragraph_style = DEFAULT_PARAGRAPH_STYLE

    def set_initial_attrs(self, document=None, headers=None, base_url=None):
        self.tags = {'span': [], 'list': []}
        self.doc = document if document else Document()
        self.headers = headers
        self.base_url = base_url
        self.bs = self.options['fix-html']
        self.document = self.doc
        self.include_tables = True
        self.include_images = self.options['images']
        self.include_styles = self.options['styles']
        self.paragraph = None
        self.skip = False
        self.skip_tag = None
        self.instances_to_skip = 0

    def copy_settings_from(self, other):
        self.table_style = other.table_style
        self.paragraph_style = other.paragraph_style

    def get_cell_html(self, soup):
        return ' '.join(str(i) for i in soup.contents)

    def add_styles_to_paragraph(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9]+', '', margin)
            margin_value = int(float(re.sub(r'[a-z]+', '', margin)))
            if units == 'px':
                self.paragraph.paragraph_format.left_indent = Inches(
                    min(margin_value // 10 * INDENT, MAX_INDENT)
                )

    def add_styles_to_run(self, style):
        if 'font-size' in style:
            px_to_pt = 1.333
            pt_size = float(style['font-size'].replace('px', '')) / px_to_pt
            self.run.font.size = Pt(pt_size)
        if 'color' in style:
            color_str = style['color']
            if 'rgb' in color_str:
                nums = re.sub(r'[a-z()]+', '', color_str)
                colors = [int(x) for x in nums.split(',')]
            elif '#' in color_str:
                hexc = color_str.lstrip('#')
                colors = tuple(int(hexc[i:i+2], 16) for i in (0, 2, 4))
            else:
                colors = (0, 0, 0)
            self.run.font.color.rgb = RGBColor(*colors)
        if 'background-color' in style:
            # Simplify: always gray if any background
            self.run.font.highlight_color = WD_COLOR.GRAY_25

    def apply_paragraph_style(self, style=None):
        try:
            if style:
                self.paragraph.style = style
            elif self.paragraph_style:
                self.paragraph.style = self.paragraph_style
        except KeyError as e:
            raise ValueError(f"Unable to apply style {self.paragraph_style}.") from e

    def parse_dict_string(self, string, separator=';'):
        parts = string.replace(" ", "").split(separator)
        return dict(x.split(':') for x in parts if ':' in x)

    def handle_li(self):
        depth = len(self.tags['list'])
        list_type = self.tags['list'][-1] if depth else 'ul'
        style = (styles['LIST_NUMBER'] if list_type == 'ol' else styles['LIST_BULLET'])
        self.paragraph = self.doc.add_paragraph(style=style)
        self.paragraph.paragraph_format.left_indent = Inches(
            min(depth * LIST_INDENT, MAX_INDENT)
        )

    def add_image_to_cell(self, cell, image):
        p = cell.add_paragraph()
        r = p.add_run()
        r.add_picture(image)

    def handle_img(self, current_attrs):
        if not self.include_images:
            self.skip = True
            self.skip_tag = 'img'
            return
        src = current_attrs.get('src')
        image = None
        try:
            image = fetch_image(src, self.headers, self.base_url)
        except Exception as e:
            print("❌ Error fetching image:", e)
        if image:
            if isinstance(self.doc, docx.document.Document):
                self.doc.add_picture(image, width=Inches(6.69291), height=Inches(4.2244094))
            else:
                self.add_image_to_cell(self.doc, image)
        else:
            self.doc.add_paragraph(f"<image: {src}>")

    # ... rest of HtmlToDocx methods remain unchanged ...

    def run_process(self, html):
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, 'html.parser')
            html = str(self.soup)
        if self.include_tables:
            self.get_tables()
        self.feed(html)

    def add_html_to_document(self, html, document, headers=None, base_url=None):
        if not isinstance(html, str):
            raise ValueError('First argument must be a str')
        self.set_initial_attrs(document, headers, base_url)
        self.run_process(html)

    def add_html_to_cell(self, html, cell):
        if not isinstance(cell, docx.table._Cell):
            raise ValueError('Second argument must be a docx cell')
        if cell.paragraphs and cell.paragraphs[0].text == "":
            delete_paragraph(cell.paragraphs[0])
        self.set_initial_attrs(cell)
        self.run_process(html)
        if not self.doc.paragraphs:
            self.doc.add_paragraph('')

    # ... other helper methods unchanged ...

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert HTML to DOCX')
    parser.add_argument('filename_html')
    parser.add_argument('filename_docx', nargs='?', default=None)
    parser.add_argument('--bs', action='store_true', help='Fix HTML via BeautifulSoup')
    args = parser.parse_args()
    html_parser = HtmlToDocx()
    html_parser.parse_html_file(args.filename_html, args.filename_docx)